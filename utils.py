from os import listdir
from os.path import isfile, join
import hashlib
import glob
import docx
import re
import string
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import difflib
import numpy as np
from sklearn.feature_extraction.text import CountVectorizer
import nltk
import pandas as pd


def get_text_name(folder):
    only_files = [f for f in listdir(folder) if isfile(join(folder, f))]
    return only_files


def get_text_dir(folder):
    file_dir = glob.glob(folder + "*")
    return file_dir


def para_text_extract(files):
    '''
        Extrai docs e parágrafos docx e textos separados em parágrafos docx e '\n'
    '''

    docs = [docx.Document(doc) for doc in files]
    texts = [doc.paragraphs for doc in docs]

    para_texts = []
    for text in texts:
        para_group = []
        for para in text:
            para_group.extend(para.text.split('\n'))
        para_texts.append(para_group)

    doc_sizes = [len(para) for para in para_texts]
    doc_sizes_total = sum(doc_sizes)

    print('Quantidade de documentos: ' + str(len(para_texts)))
    print('Tamanho de cada documento: ' + str(doc_sizes))
    print('Tamanho total: ' + str(doc_sizes_total))

    return docs, texts, para_texts

def get_hash(file_dir, chunk_size=1024):
    '''
        Recebe o caminho de um arquivo e retorna o hash acumulado
    '''
    hash = hashlib.sha256()
    with open(file_dir, "rb") as f:
        # lê o primeiro bloco do arquivo
        chunk = f.read(chunk_size)
        # fica lendo arquivo até o fim e atualiza o hash
        while chunk:
            hash.update(chunk)
            chunk = f.read(chunk_size)

    # Returna hex checksum
    return hash.hexdigest()



def date_extract(docs):
    '''
        Extrai datas dos cabeçalhos
    '''

    dates = []
    for doc in docs:
        section = doc.sections[0]
        header = section.first_page_header
        try:
            date = header.paragraphs[1].text
        except IndexError:
            date = None
        dates.append(date)

    return dates


def full_text_extract(filename):
    '''
        Extrai texto completo
    '''

    doc = docx.Document(filename)

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    return '\n'.join(full_text)


def combine_text(list_of_text):
    '''
        Une as strings de texto com caractere espaço
    '''

    combined_text = ' '.join(list_of_text)
    return combined_text


def clean_text_round1(text):
    '''
        Eliminaa pontuação, colchetes e palavras que contribuam pouco
    '''

    text = text.lower()
    text = re.sub('\[.*?\]','',text)
    text = re.sub('[%s]' %re.escape(string.punctuation
                                    ),'',text)
    text = re.sub('\w*\d\w*','',text)
    return text


def clean_text_round2(text):
    '''
        Segunda rodada de limpeza
    '''
    
    text = re.sub('[‘’“”…]', '', text)
    text = re.sub('\n', '', text)
    text = re.sub('\d{2}.\d{3}.\d{3}/\d{4}-\d{2}','"CNPJ"',text)
    text = re.sub('\(s\)','',text)
    text = re.sub('\(\w{2}\)','',text)
    return text


def similarities(clauses, text):
    '''
        Calcula a similaridade entre uma lista de cláusulas e cada segmento de texto e
        retorna uma lista contendo listas com as razões de similaridade de cada cláusula
    Args:
        clauses - lista de strings contendo cláusulas obrigatórias
        text - lista de parágrafos de um texto
    Return:
        Lista contendo listas com os índices de similaridade de cada clásula
    '''

    similarity_indices = []
    for clause in clauses:
        similarity_group = []
        for para in text:
            ratio = difflib.SequenceMatcher(None, clause, para).ratio()
            similarity_group.append(ratio)
        similarity_indices.append(similarity_group)
    return np.array(similarity_indices)


def clauses_in_text(similarities_indices):
    '''
        Veririca se uma lista de cláusulas está no texto pela razão de similaridade >= 0.8
    Arg:
        similarities_indices - Lista contendo listas com os índices de similaridade de cada clásula
    Return:
        Quantidade de ocorrências das cláusulas no texto
    '''

    mask = similarities_indices >= 0.8
    lines = np.any(mask, axis=0)
    if np.any(lines==True):
        return sum(lines)
    else:
        return 0


def included_clauses(similarities_indices):
    '''
        Gera uma lista indicando se as cláusulas estão presentes em cada parágrafo do texto
    Arg:
        similarity_indices - Lista com os índices de similaridade da clásula em cada parágrago do texto
    Return:
        Lista com 0s e 1s indicando se a cláusula está presente em cada parágrafo do texto.
    '''

    mask = similarities_indices >= 0.8
    lines = np.any(mask, axis=0)
    return lines


def create_dtm(docs_df):
    '''
        Essa função cria o DTM de todos os textos do corpus.
        Tem como parametro de entrada um dataframe que representa o corpus de textos
    '''

    stopwords = nltk.corpus.stopwords.words('portuguese')
    cv = CountVectorizer(stop_words=stopwords)
    data_cv = cv.fit_transform(docs_df.texto)
    data_dtm = pd.DataFrame(data_cv.toarray(), columns=cv.get_feature_names_out())
    data_dtm.index = docs_df.documento
    return cv, data_dtm


def separar_parrafos_de_clausulas(df2):
    ''' 
        Pega um dataframe com os documentos separados por cláusulas
        e logo divide ele em paragrafos estiquetados por cláusulas
     '''

    Data_Clausulas = pd.DataFrame({'PARAGRAFOS':0,'CLÁUSULA':0},index=[])
    for index in df2.index:
        para_texts = []
        df = pd.DataFrame(df2.iloc[index]).transpose()
        for texts in df['textos']:
            para = texts.split('\n')
            #print(para)
            for p in para:
                #print(p)
                if len(p) == 0:
                    continue
                para_texts.append(p)
        df_cla = pd.DataFrame(para_texts,columns=['PARAGRAFOS'])
        df_cla
        #print(df.clausulas.to_list()[0])
        df_cla['CLÁUSULA'] = df.clausulas.to_list()[0]
        Data_Clausulas = Data_Clausulas.append(df_cla)
        Data_Clausulas.to_pickle('Data_Clausulas_por_paragrafos.pkl')
    return Data_Clausulas
